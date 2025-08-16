import os
import re
import fitz  # PyMuPDF
import docx
import spacy
from transformers import pipeline
from PIL import Image
from io import BytesIO
import pytesseract
import argparse

# --- 1. SETUP & CONFIGURATION ---

# Keywords for detailed checks
OUTDATED_TECH = {'flash', 'jquery', 'svn', 'subversion', 'visual basic', 'vb.net', 'webforms', 'soap', 'angularjs'}
BUZZWORDS = {'synergy', 'go-getter', 'team player', 'results-oriented', 'dynamic', 'proactive', 'thought leader', 'self-starter'}
ACTION_VERBS = {'achieved', 'accelerated', 'improved', 'drove', 'managed', 'created', 'launched', 'led', 'increased', 'decreased', 'optimized'}
EDUCATION_TERMS = {'b.s', 'b.a', 'm.s', 'm.a', 'ph.d', 'bachelor', 'master', 'doctorate', 'university', 'college', 'institute'}

# Domain keywords for smarter inference
DOMAIN_KEYWORDS = {
    "software": {"engineer", "developer", "python", "java", "aws", "software", "api", "database", "git", "cloud"},
    "marketing": {"marketing", "campaign", "brand", "seo", "ppc", "content", "digital", "analytic", "social media"},
    "graphic_designer": {"graphic", "photoshop", "illustrator", "figma", "design", "visual", "assets", "branding"},
    "scm": {"supply chain", "scm", "logistics", "warehousing", "distribution", "procurement", "inventory", "freight"}
}

DEFAULT_JDS = {
    "marketing": "Marketing Manager: Expertise in B2B/B2C marketing, demand generation, campaign execution, digital marketing (SEO, SEM, PPC), content strategy, and marketing analytics.",
    "software": "Senior Software Engineer: Skilled in Python, Java, Go, data structures, algorithms, cloud computing (AWS/Azure/GCP), microservices architecture, Docker, Kubernetes, and SQL/NoSQL databases.",
    "graphic_designer": "Graphic Designer: Proficiency in Adobe Creative Suite (Photoshop, Illustrator, InDesign), Figma, and Sketch. Experience in UI/UX design, branding, and creating visual assets for digital and print media.",
    "scm": "Supply Chain Manager: Experience in logistics, inventory management, procurement, vendor relations, warehousing, and distribution. Focus on optimizing processes and reducing costs.",
    "generic": "A general professional role with a focus on clear communication, quantifiable achievements, leadership, project management, and problem-solving skills."
}

# --- Graceful Model Loading ---
try:
    nlp = spacy.load("en_core_web_sm")
except Exception as e:
    print(f"[ERROR] Failed to load spaCy model. Keyword analysis will be disabled. Error: {e}")
    nlp = None

try:
    summarizer = pipeline("summarization", model="sshleifer/distilbart-cnn-12-6")
except Exception as e:
    print(f"[WARN] Could not load summarization model. AI summaries will be unavailable. Error: {e}")
    summarizer = None

# --- 2. TEXT EXTRACTION & PARSING ENGINE ---

def extract_text_from_pdf(filepath):
    """Extracts text from a PDF, automatically applying OCR if needed."""
    full_text, is_graphics_heavy = "", False
    try:
        doc = fitz.open(filepath)
        for page in doc:
            full_text += page.get_text()
        if len(full_text.strip()) < 150:
            is_graphics_heavy = True
            print("[INFO] Low text content detected. Attempting OCR...")
            full_text = ""
            for page in doc:
                pix = page.get_pixmap(dpi=300)
                img = Image.open(BytesIO(pix.tobytes()))
                try:
                    page_text = pytesseract.image_to_string(img, lang='eng')
                    full_text += page_text + "\n"
                except pytesseract.TesseractNotFoundError:
                    return "TESSERACT_NOT_FOUND", False
            print("[INFO] OCR processing complete.")
        return full_text.strip(), is_graphics_heavy
    except Exception as e:
        print(f"[ERROR] Failed to process PDF {filepath}: {e}")
        return "", False

def extract_text_from_docx(filepath):
    try:
        doc = docx.Document(filepath)
        return "\n".join([para.text for para in doc.paragraphs]).strip(), False
    except Exception as e:
        print(f"[ERROR] Failed to read .docx file {filepath}: {e}")
        return "", False

def parse_sections_by_iteration(text):
    """
    Parses the resume text line-by-line to accurately identify sections.
    """
    section_keywords = {
        "profile_summary": ["summary", "profile", "objective"],
        "work_experience": ["experience", "work history", "professional experience", "experience details"],
        "education": ["education", "academic background", "academic record", "professional qualification"],
        "skills": ["skills", "technical skills", "technical proficiency", "tools", "skill set", "core competencies", "areas of expertise", "computer skills"]
    }
    
    keyword_to_section = {kw: key for key, kws in section_keywords.items() for kw in kws}
    
    sections = {key: [] for key in section_keywords.keys()}
    current_section = None
    
    lines = text.split('\n')
    for line in lines:
        cleaned_line = line.strip().lower().replace(':', '')
        
        if 1 <= len(cleaned_line.split()) <= 4 and cleaned_line in keyword_to_section:
            current_section = keyword_to_section[cleaned_line]
            continue

        if current_section:
            sections[current_section].append(line)
            
    for section_name, section_lines in sections.items():
        sections[section_name] = "\n".join(section_lines).strip()
        
    return sections

def master_text_extractor(filepath):
    ext = os.path.splitext(filepath)[1].lower()
    raw_text, is_graphics_heavy = "", False
    if ext == ".pdf":
        raw_text, is_graphics_heavy = extract_text_from_pdf(filepath)
    elif ext == ".docx":
        raw_text, is_graphics_heavy = extract_text_from_docx(filepath)
    else:
        return None, {}, False
    if not raw_text or raw_text == "TESSERACT_NOT_FOUND":
        return raw_text, {}, is_graphics_heavy
    
    normalized_text = re.sub(r'[ \t]+', ' ', raw_text)
    normalized_text = re.sub(r'\s*\n\s*', '\n', normalized_text).strip()
    
    sections = parse_sections_by_iteration(normalized_text)
    return normalized_text, sections, is_graphics_heavy

# --- 3. ADVANCED ANALYSIS ENGINE ---

def generate_ai_summary(resume_text):
    if not summarizer: return "AI summary model is unavailable."
    try:
        cleaned_text = re.sub(r'\s*\n\s*', ' ', resume_text).strip()
        summary = summarizer(cleaned_text[:2048], max_length=130, min_length=40, do_sample=False)
        return summary[0]['summary_text'].strip()
    except Exception: return "Could not generate a unique AI summary."

def match_job_description(resume_text, job_description_text):
    if not nlp: return {"score": 0, "missing_keywords": [], "overlap_keywords": []}
    resume_lemmas = {token.lemma_ for token in nlp(resume_text.lower()) if not token.is_stop and not token.is_punct and token.is_alpha}
    jd_lemmas = {token.lemma_ for token in nlp(job_description_text.lower()) if not token.is_stop and not token.is_punct and token.is_alpha}
    if not jd_lemmas: return {"score": 0, "missing_keywords": [], "overlap_keywords": []}
    overlap = resume_lemmas & jd_lemmas
    score = int((len(overlap) / len(jd_lemmas)) * 100) if jd_lemmas else 0
    return {
        "score": score, 
        "missing_keywords": sorted(list(jd_lemmas - resume_lemmas)),
        "overlap_keywords": sorted(list(overlap))
    }

def compile_issues(sections, raw_text):
    formatting_issues, content_issues = [], []
    
    if not sections.get("skills"): formatting_issues.append("A dedicated 'Skills' section is crucial for ATS and is missing.")
    if not sections.get("work_experience"): formatting_issues.append("A standard 'Work Experience' section was not found.")
    
    summary_text = sections.get("profile_summary", "")
    if summary_text:
        found_buzzwords = BUZZWORDS.intersection(summary_text.lower().split())
        if found_buzzwords: content_issues.append(f"Profile summary uses buzzwords: {', '.join(found_buzzwords)}.")
        if not any(char.isdigit() for char in summary_text) and not any(verb in summary_text.lower() for verb in ACTION_VERBS):
            content_issues.append("Profile summary may lack a clear value proposition.")

    education_text = sections.get("education", "")
    if education_text:
        if not any(term in education_text.lower() for term in EDUCATION_TERMS): content_issues.append("Education section may be missing degree or institution information.")
        if not re.search(r'\b(19|20)\d{2}\b', education_text): content_issues.append("Education section is missing a clear graduation year.")
    else:
        content_issues.append("An 'Education' section was not found.")

    exp_text = sections.get("work_experience", "")
    if exp_text:
        lines = [line for line in exp_text.split('\n') if len(line.strip()) > 15]
        quantified_lines = sum(1 for line in lines if any(char.isdigit() for char in line))
        quantification_ratio = (quantified_lines / len(lines)) if lines else 0
        if quantification_ratio < 0.3: content_issues.append("Work experience lacks quantification. Add metrics to show impact.")
        
    outdated_tech = list(OUTDATED_TECH.intersection(token.lemma_.lower() for token in nlp(raw_text))) if nlp else []

    return formatting_issues, content_issues, outdated_tech

def calculate_scores(formatting_issues, content_issues, outdated_tech, jd_match_score):
    total_failures = len(formatting_issues) + len(content_issues)
    quality_score = max(0, 100 - (total_failures * 10) - (len(outdated_tech) * 5))
    final_score = round((quality_score * 0.4) + (jd_match_score * 0.6))
    grade = next((g for t, g in {80: "A", 70: "B", 60: "C", 40: "D"}.items() if final_score >= t), "F")
    return final_score, grade

def generate_personalized_recommendations(report):
    recommendations, checks = [], report['classified_checks']
    job_fit = checks.get("Job Fit", {})
    if job_fit.get("score", 100) < 75 and job_fit.get("missing_keywords"):
        recommendations.append(f"**Improve Job Fit:** Your resume is missing key terms from the job description. To better match the role, integrate these keywords: **{', '.join(job_fit['missing_keywords'][:8])}**.")
    ats_issues = checks.get("Formatting & ATS", {}).get("issues", [])
    if any("section" in issue for issue in ats_issues):
        recommendations.append("**Check Section Headers:** To ensure automated systems read your resume correctly, use standard titles like **'Work Experience'** and **'Technical Skills'**.")
    content_issues = checks.get("Content & Impact", {}).get("issues", [])
    if any("value proposition" in issue for issue in content_issues):
        recommendations.append("**Rewrite Your Profile Summary:** Your summary may use generic phrases. Create a stronger value proposition with powerful action verbs and achievements.")
    if any("quantification" in issue for issue in content_issues):
         recommendations.append("**Quantify Your Achievements:** Your work experience lacks metrics. Strengthen your bullet points by adding numbers, percentages (%), or dollar amounts ($).")
    if checks.get("Technology Freshness", {}).get("outdated_tech"):
        recommendations.append(f"**Update Your Technology Stack:** Your resume mentions potentially outdated technologies. Prioritize highlighting modern tools relevant to the job.")
    if not recommendations:
        recommendations.append("Your resume is strong and well-aligned. No major recommendations at this time.")
    return recommendations

# --- 4. MAIN ORCHESTRATOR & REPORTING ---

def analyze_resume(filepath, job_description_text=None):
    raw_text, sections, is_graphics_heavy = master_text_extractor(filepath)
    if not raw_text: return {"error": "Failed to extract text. File may be corrupted or an unsupported format."}
    if raw_text == "TESSERACT_NOT_FOUND": return {"error": "Tesseract OCR is not installed."}

    if not job_description_text:
        print("[INFO] No job description provided. Inferring from resume content...")
        text_lower = raw_text.lower()
        domain_scores = {domain: sum(1 for keyword in keywords if keyword in text_lower) for domain, keywords in DOMAIN_KEYWORDS.items()}
        best_domain = max(domain_scores, key=domain_scores.get)
        domain = best_domain if domain_scores[best_domain] > 2 else "generic"
        job_description_text = DEFAULT_JDS.get(domain)
        print(f"[INFO] Comparing against the default '{domain}' profile.")

    jd_match = match_job_description(raw_text, job_description_text)
    formatting_issues, content_issues, outdated_tech = compile_issues(sections, raw_text)
    if is_graphics_heavy: formatting_issues.append("Resume appears to be graphics-heavy or image-based.")
    
    final_score, grade = calculate_scores(formatting_issues, content_issues, outdated_tech, jd_match["score"])
    
    report = {
        "overall_score": final_score, "resume_grade": grade,
        "ai_generated_summary": generate_ai_summary(raw_text),
        "classified_checks": {
            "Formatting & ATS": {"issues": formatting_issues},
            "Content & Impact": {"issues": content_issues},
            # --- UPDATED: Added overlap_keywords to the report ---
            "Job Fit": {
                "score": jd_match['score'], 
                "missing_keywords": jd_match['missing_keywords'],
                "overlap_keywords": jd_match['overlap_keywords'] 
            },
            "Technology Freshness": {"outdated_tech": outdated_tech}
        }
    }
    report["personalized_recommendations"] = generate_personalized_recommendations(report)
    return report

# This part is for direct execution, which we are replacing with a web server.
# You can keep it for testing purposes if you like.
if __name__ == '__main__':
    parser = argparse.ArgumentParser(description="Analyze a resume against a job description.")
    parser.add_argument("resume_path", help="The full path to the resume file (PDF or DOCX).")
    parser.add_argument("--jd", dest="jd_path", help="Optional path to a text file containing the job description.")
    
    args = parser.parse_args()
    
    if not os.path.exists(args.resume_path):
        print(f"[ERROR] Resume file not found at: {args.resume_path}")
    else:
        job_description = None
        if args.jd_path:
            if os.path.exists(args.jd_path):
                with open(args.jd_path, 'r', encoding='utf-8') as f:
                    job_description = f.read()
            else:
                print(f"[WARN] Job description file not found at: {args.jd_path}. Using default profile.")

        print(f"🚀 Analyzing '{os.path.basename(args.resume_path)}'...")
        final_report = analyze_resume(args.resume_path, job_description_text=job_description)
        
        if "error" in final_report:
            print(f"Error: {final_report['error']}")
        else:
            print("\n" + "="*25 + " ANALYSIS COMPLETE " + "="*25)
            print(f"Overall Score: {final_report['overall_score']}/100 (Grade: {final_report['resume_grade']})")
            
            print("\n💡 PERSONALIZED ACTION PLAN")
            print("-" * 50)
            for rec in final_report['personalized_recommendations']:
                rec_text = rec.replace("**", "").replace("*", "")
                print(f"- {rec_text}")
            
            print("\n🔬 DETAILED ISSUES FOUND")
            print("-" * 50)
            for category, checks in final_report['classified_checks'].items():
                all_issues = checks.get('issues', [])
                if checks.get('outdated_tech'):
                    all_issues.append(f"Mentions outdated tech: {', '.join(checks.get('outdated_tech'))}")
                
                if all_issues:
                    print(f"[{category}]")
                    for issue in all_issues:
                        print(f"  - {issue}")

            print("\n" + "="*69)
if __name__ == '__main__':
    try:
        doc = docx.Document()
        doc.add_heading("Profile", level=1)
        doc.add_paragraph("A dynamic team player looking for a software engineer role.")
        doc.save("dummy_resume_no_jd.docx")
        
        resume_file_path = "dummy_resume_no_jd.docx"
        
        # --- Run without a specific job description ---
        print("🚀 Analyzing resume with NO job description (should use default)...")
        final_report = analyze_resume(resume_file_path) # Calling with only the filepath
        
        if "error" in final_report:
            print(f"Error: {final_report['error']}")
        else:
            print("\n" + "="*50)
            print("💡 PERSONALIZED ACTION PLAN (from default JD) 💡")
            print("="*50)
            for i, rec in enumerate(final_report['personalized_recommendations'], 1):
                rec_text = rec.replace("**", "").replace("*", "")
                print(f"{i}. {rec_text}")
            print(f"\nJob Fit Score (against default): {final_report['classified_checks']['Job Fit']['score']}%")
            print("="*50 + "\n")

    except Exception as e:
        print(f"An error occurred during the example run: {e}")
    finally:
        if os.path.exists("dummy_resume_no_jd.docx"):
            os.remove("dummy_resume_no_jd.docx")